# -*- coding: utf-8 -*-
"""DB → DataFrame → Automatic Experimental Description → Word(docx) + Write to report_edge.

Supports batch generation for multiple tables.
- Deduplication via unique index
- Status codes: 0=skipped, 1=generated, 2=failed
"""

import os
import random
from datetime import datetime
from pathlib import Path

import pandas as pd
import pymysql
from docx import Document
from sqlalchemy import create_engine

from seven_ai_layers_robotics.config import config
from templates_new_revised import (
    get_intro_segment,
    perovskite_formula_segments,
    sam_formula_segments_single,
    sam_formula_segments_dual,
    sam_formula_segments_triple,
    sam_formula_spin_segments_single,
    sam_formula_spin_segments_dual,
    sam_formula_spin_segments_triple,
    sam_spin_anneal_segments,
    additive_formula_segments_single,
    additive_formula_segments_dual,
    additive_formula_segments_triple,
    passivation_material_segments_single,
    passivation_material_segments_dual,
    passivation_material_segments_triple,
    process_segments,
    antisolvent_segments,
    anneal_segments,
    passivation_spin_segments,
    passivation_drop_segments,
    passivation_anneal_segments,
    image_analysis_segments,
    pl_analysis_segments,
    xrd_analysis_segments_12,
    xrd_analysis_segments_stress,
)

MYSQL_CONFIG = {
    'host': config.generating_database.host,
    'port': config.generating_database.port,
    'user': config.generating_database.user,
    'password': config.generating_database.password,
    'database': config.generating_database.database,
    'charset': config.generating_database.charset,
}
DB_URI = f"mysql+pymysql://{MYSQL_CONFIG['user']}:{MYSQL_CONFIG['password']}@{MYSQL_CONFIG['host']}:{MYSQL_CONFIG['port']}/{MYSQL_CONFIG['database']}?charset=utf8mb4"
engine = create_engine(DB_URI)
DB_CONFIG = MYSQL_CONFIG


def choose_random(lst):
    """Select a random element from a list.
    
    Args:
        lst: A non-empty list to select from.
        
    Returns:
        A randomly selected element from the list.
    """
    return random.choice(lst)

def format_number_int_like(x, ndigits=2):
    """Format a number to integer-like string representation.
    
    Args:
        x: Input number to format.
        ndigits: Number of decimal places for rounding comparison. Defaults to 2.
        
    Returns:
        String representation (integer if close to whole number, otherwise formatted decimal).
    """
    if pd.isna(x):
        return ""
    try:
        x = float(x)
    except Exception:
        return str(x)
    if abs(x - round(x)) < 10 ** -ndigits:
        return str(int(round(x)))
    return f"{x:.{ndigits}f}".rstrip("0").rstrip(".")


def format_multiple_materials(materials, label) -> str:
    """Format multiple material names and concentrations into a descriptive string.
    
    Args:
        materials: List of tuples containing (formula, concentration) pairs.
        label: Label for the material group (e.g., 'additives', 'SAMs').
        
    Returns:
        Formatted descriptive string of materials.
    """
    parts = []
    for f, c in materials:
        if pd.isna(f) and pd.isna(c):
            continue
        c = "" if pd.isna(c) else str(c)
        if c and not any(ch.isalpha() for ch in c):
            c += " mg/mL"
        parts.append(f"{f} ({c})" if c else f"{f}")
    return f"The {label} included " + ", ".join(parts) + "." if parts else ""


def read_table(table) -> pd.DataFrame:
    """Read a database table into a pandas DataFrame.
    
    Args:
        table: Name of the database table to read.
        
    Returns:
        pandas DataFrame containing the table data with cleaned column names.
    """
    df = pd.read_sql(f"SELECT * FROM `{table}`;", con=engine)
    df.columns = [str(c).strip() for c in df.columns]
    return df


def safe_float(x, ndigits=None):
    """Safely convert a value to float with optional rounding.
    
    Args:
        x: Value to convert to float.
        ndigits: Number of decimal places for rounding. If None, no rounding.
        
    Returns:
        Float value or None if conversion fails.
    """
    try:
        x = float(x)
        return round(x, ndigits) if ndigits is not None else x
    except Exception:
        return None


def generate_paragraph(row, xrd_mode=None) -> str:
    """Generate experimental description paragraph from a data row.
    
    Args:
        row: Dictionary-like row containing experimental data.
        xrd_mode: XRD analysis mode ('additives', 'passivators', or None).
        
    Returns:
        Generated paragraph text describing the experiment.
    """
    pce = safe_float(row.get('PCE'), 4)
    ff  = safe_float(row.get('FF'), 4)
    voc = safe_float(row.get('Voc'), 4)
    jsc = safe_float(row.get('Jsc'), 3)

    if any(v is None for v in [pce, ff, voc, jsc]):
        return ""

    intro = get_intro_segment(pce=pce, ff=ff, voc=voc, jsc=jsc)
    perovskite = choose_random(perovskite_formula_segments).format(
        formula_pvk=row['Formula PVK'], concentration_pvk=row['Concentration PVK']
    )

    process = choose_random(process_segments).format(
        spin1_speed=row['Spin Coating Speed PVK 1'],
        spin1_time=row['Spin Coating Time PVK 1'],
        spin2_speed=row['Spin Coating Speed PVK 2'],
        spin2_time=row['Spin Coating Time PVK 2']
    )

    antisolvent = choose_random(antisolvent_segments).format(
        antisolvent_volume=row['Antisolvent Volume'],
        antisolvent_timing=row['Antisolvent Dropping Timing']
    )

    anneal = choose_random(anneal_segments).format(
        anneal_temp=row['Annealed Temperature PVK'],
        anneal_time=row['Annealed Time PVK']
    )
    parts = [intro, perovskite, process, antisolvent, anneal]

    def has_any(keys):
        return any(k in row and pd.notna(row[k]) and str(row[k]).strip() for k in keys)
    if has_any(["area_px2", "gray_mean"]):
        for tpl in image_analysis_segments:
            parts.append(tpl.format(area_px2=row.get("area_px2",""), gray_mean=row.get("gray_mean","")))

    if has_any(["peak_time", "decay_slope"]):
        for tpl in pl_analysis_segments:
            parts.append(tpl.format(
                peak_time=row.get("peak_time",""),
                decay_slope=row.get("decay_slope","")
            ))

    if xrd_mode == "additives" and has_any(["xrd_intensity_12.6"]):
        for tpl in xrd_analysis_segments_12:
            parts.append(tpl.format(
                xrd_intensity_12=f"{float(row['xrd_intensity_12.6']):.2f}",
                xrd_fhwm_12=f"{float(row['xrd_fhwm_12.6']):.2f}"
            ))
    if xrd_mode == "passivators" and has_any(["xrd_intensity_4"]):
        for tpl in xrd_analysis_segments_stress:
            parts.append(tpl.format(
                xrd_intensity_4=f"{float(row['xrd_intensity_4']):.2f}",
                xrd_Stress=f"{float(row['xrd_Stress']):.2f}"
            ))

    return " ".join(p for p in parts if p)


def insert_single_report_records(table_name: str, record_ids: list, status_code: int, location: str) -> None:
    """Insert single report records into database with deduplication.
    
    Args:
        table_name: Source table name for reference.
        record_ids: List of record IDs to insert.
        status_code: Status code (0=skipped, 1=generated, 2=failed).
        location: File path location of the generated report.
        
    Returns:
        None
    """
    if not record_ids:
        return
    conn = pymysql.connect(**DB_CONFIG, connect_timeout=30)
    cursor = conn.cursor()

    uploadtime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    sql = """
    INSERT IGNORE INTO `report_edge` (`type`, `id`, `status`, `uploadtime`, `location`)
    VALUES (%s, %s, %s, %s, %s)
    """
    data = [(table_name, rid, status_code, uploadtime, location) for rid in record_ids]
    cursor.executemany(sql, data)
    inserted_count = cursor.rowcount
    conn.commit()
    cursor.close()
    conn.close()

    if inserted_count > 0:
        status_map = {0: "skipped", 1: "generated", 2: "failed"}
        print(f"Added {inserted_count} records to report_edge (status={status_map[status_code]})")


def ensure_single_report_unique_index() -> None:
    """Ensure unique index exists on report_edge table for (type, id) combination.
    
    Creates the unique index if it doesn't exist to prevent duplicate entries.
    
    Returns:
        None
    """
    conn = pymysql.connect(**DB_CONFIG, connect_timeout=30)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT COUNT(1) FROM information_schema.statistics
        WHERE table_schema = %s
          AND table_name = 'report_edge'
          AND index_name = 'uniq_type_id';
    """, (DB_CONFIG['database'],))
    exists = cursor.fetchone()[0] > 0
    if not exists:
        cursor.execute("ALTER TABLE `report_edge` ADD UNIQUE KEY `uniq_type_id` (`type`, `id`);")
        print("Added unique index (type, id) to report_edge")
    cursor.close()
    conn.close()


def generate_docx_from_df(df, output_path, xrd_mode=None, index_col=None, source_table=None) -> None:
    """Generate Word document from DataFrame and record status to database.
    
    Args:
        df: pandas DataFrame containing experimental data.
        output_path: Output file path for the generated Word document.
        xrd_mode: XRD analysis mode ('additives', 'passivators', or None).
        index_col: Column name to use as record ID. If None, uses row index.
        source_table: Source database table name for recording status.
        
    Returns:
        None
        
    Side Effects:
        - Creates Word document at output_path
        - Inserts records into report_edge table with status codes
    """
    doc = Document()
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    generated_ids = []
    skipped_ids = []
    failed_ids = []

    if not index_col or index_col not in df.columns:
        index_col = None

    for i, row in df.iterrows():
        if index_col:
            raw_id = row.get(index_col)
        else:
            raw_id = i + 1

        try:
            if pd.isna(raw_id):
                record_id = i + 1
            elif isinstance(raw_id, float) and raw_id.is_integer():
                record_id = int(raw_id)
            else:
                record_id = raw_id
        except Exception:
            record_id = str(raw_id)

        try:
            text = generate_paragraph(row, xrd_mode=xrd_mode)
            if not text or not text.strip():
                skipped_ids.append(record_id)
            else:
                idx_show = record_id
                doc.add_paragraph(f"{idx_show}. {text}")
                doc.add_paragraph("")
                generated_ids.append(record_id)
        except Exception as e:
            print(f"⚠️ Generation failed (ID={record_id}): {e}")
            failed_ids.append(record_id)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    total = len(generated_ids) + len(skipped_ids) + len(failed_ids)
    print(f"Completed {source_table}: total {total} | generated {len(generated_ids)} | skipped {len(skipped_ids)} | failed {len(failed_ids)}")

    if source_table:
        insert_single_report_records(source_table, generated_ids, 1, output_path)
        insert_single_report_records(source_table, skipped_ids, 0, output_path)
        insert_single_report_records(source_table, failed_ids, 2, output_path)

# =========================
# Batch Task Configuration
# =========================
TASKS = [
    {"table": "characterisation_xrd_additives", "output": "reports/characterisation_xrd_additives.docx", "xrd_mode": "additives", "index_col": "index"},
    {"table": "characterisation_xrd_passivators", "output": "reports/characterisation_xrd_passivators.docx", "xrd_mode": "passivators", "index_col": "index"},
    {"table": "characterisation_image_pvk", "output": "reports/characterisation_image_pvk.docx", "xrd_mode": None, "index_col": "index"},
    {"table": "characterisation_pl_sam", "output": "reports/characterisation_pl_sam.docx", "xrd_mode": None, "index_col": "index"},
    {"table": "experiments_cleaned_data", "output": "reports/experiments_cleaned_data.docx", "xrd_mode": None, "index_col": "No"},
]


def main() -> None:
    """Main entry point for edge report generation."""
    ensure_single_report_unique_index()

    for task in TASKS:
        print(f"\nProcessing table: {task['table']}")
        df = read_table(task["table"])
        abs_output = os.path.abspath(task["output"])
        generate_docx_from_df(
            df=df,
            output_path=abs_output,
            xrd_mode=task["xrd_mode"],
            index_col=task.get("index_col"),
            source_table=task["table"]
        )
    print("\nAll tasks completed!")


if __name__ == "__main__":
    main()
